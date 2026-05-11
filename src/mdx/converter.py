from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token

from .styles import StylePreset, apply_document_style, rgb_hex


@dataclass(frozen=True)
class ConvertResult:
    input_path: Path
    output_path: Path
    style: str


@dataclass(frozen=True)
class ReportHeader:
    title: str
    metadata: tuple[tuple[str, str], ...]
    body_markdown: str


def convert_markdown_to_docx(
    input_path: Path,
    output_path: Path,
    *,
    style: str,
    template_path: Path | None = None,
) -> ConvertResult:
    markdown = input_path.read_text(encoding="utf-8")
    document = Document(template_path) if template_path else Document()
    preset = apply_document_style(document, style)
    header = _extract_report_header(markdown)
    title = header.title or input_path.stem.replace("-", " ").title()
    document.core_properties.title = title
    _apply_report_header_footer(document, title, preset)
    if header.metadata:
        _add_cover_page(document, header, preset)
        markdown = header.body_markdown
    _render_markdown(document, markdown, preset)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return ConvertResult(input_path=input_path, output_path=output_path, style=style)


def render_markdown_to_html(markdown: str) -> str:
    parser = MarkdownIt("commonmark", {"breaks": True}).enable("table")
    body = parser.render(markdown)
    return body


def _render_markdown(document: DocumentType, markdown: str, preset: StylePreset) -> None:
    parser = MarkdownIt("commonmark", {"breaks": True}).enable("table")
    tokens = parser.parse(markdown)

    i = 0
    # Each entry: ("bullet"|"ordered", start_value)
    list_stack: list[tuple[str, int]] = []
    blockquote_depth = 0

    while i < len(tokens):
        token = tokens[i]

        if token.type == "heading_open":
            level = int(token.tag[1])
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            paragraph = document.add_heading(level=level)
            # H1 and H2 get a subtle bottom rule for visual separation.
            if level == 1:
                _set_paragraph_bottom_border(paragraph, color=rgb_hex(preset.accent_color), size="10")
            elif level == 2:
                _set_paragraph_bottom_border(paragraph, color="D9D9D9", size="6")
            if inline and inline.type == "inline":
                _append_inline(paragraph, inline, preset)
            i += 3
            continue

        if token.type == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            paragraph = document.add_paragraph()
            if list_stack:
                list_kind, _ = list_stack[-1]
                depth = len(list_stack)
                clamped = min(depth, 3)
                if list_kind == "bullet":
                    style_name = "List Bullet" if depth == 1 else f"List Bullet {clamped}"
                else:
                    style_name = "List Number" if depth == 1 else f"List Number {clamped}"
                paragraph.style = style_name
            elif blockquote_depth:
                paragraph.style = "Quote"
                _set_paragraph_left_border(paragraph, color=rgb_hex(preset.accent_color))
            if inline and inline.type == "inline":
                _append_inline(paragraph, inline, preset)
            i += 3
            continue

        if token.type == "bullet_list_open":
            list_stack.append(("bullet", 0))
            i += 1
            continue

        if token.type == "ordered_list_open":
            start_value = int(token.attrGet("start") or 1)
            list_stack.append(("ordered", start_value))
            i += 1
            continue

        if token.type in {"bullet_list_close", "ordered_list_close"}:
            if list_stack:
                list_stack.pop()
            i += 1
            continue

        if token.type in {"fence", "code_block"}:
            paragraph = document.add_paragraph(style="MDX Code")
            _shade_paragraph(paragraph, preset.code_fill)
            run = paragraph.add_run()
            _append_preserved_text(run, token.content.rstrip("\n"))
            run.font.name = preset.code_font
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), preset.code_font)
            run.font.size = Pt(max(9, preset.body_size - 1))
            i += 1
            continue

        if token.type == "blockquote_open":
            blockquote_depth += 1
            i += 1
            continue

        if token.type == "blockquote_close":
            blockquote_depth = max(0, blockquote_depth - 1)
            i += 1
            continue

        if token.type == "hr":
            paragraph = document.add_paragraph()
            _set_paragraph_bottom_border(paragraph, color=rgb_hex(preset.accent_color), size="8")
            i += 1
            continue

        if token.type == "table_open":
            i = _render_table(document, tokens, i, preset)
            continue

        i += 1


def _append_inline(paragraph, inline: Token, preset: StylePreset) -> None:
    """Render all inline tokens from *inline* into runs on *paragraph*.

    Handles: text, bold, italic, strikethrough, inline code (with background
    shading), links (underlined, colored), soft/hard breaks, and images
    (rendered as a bracketed placeholder so the document stays readable).
    """
    bold = 0
    italic = 0
    strike = 0
    link_target: str | None = None

    for child in inline.children or []:
        if child.type == "text":
            run = paragraph.add_run(child.content)
            run.bold = bold > 0
            run.italic = italic > 0
            run.font.strike = strike > 0
            if link_target is not None:
                run.underline = True
                run.font.color.rgb = RGBColor(5, 99, 193)

        elif child.type == "code_inline":
            run = paragraph.add_run(child.content)
            run.font.name = preset.code_font
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), preset.code_font)
            run.font.size = Pt(max(9, preset.body_size - 1))
            _apply_run_shading(run, preset.code_fill)

        elif child.type == "image":
            # Inline images can't be embedded without the source file; render
            # a descriptive placeholder so content is not silently lost.
            alt = child.attrGet("alt") or "image"
            run = paragraph.add_run(f"[Image: {alt}]")
            run.italic = True
            run.font.color.rgb = RGBColor(120, 120, 120)

        elif child.type == "softbreak":
            paragraph.add_run().add_break()

        elif child.type == "hardbreak":
            paragraph.add_run().add_break(WD_BREAK.LINE)

        elif child.type == "strong_open":
            bold += 1
        elif child.type == "strong_close":
            bold = max(0, bold - 1)

        elif child.type == "em_open":
            italic += 1
        elif child.type == "em_close":
            italic = max(0, italic - 1)

        elif child.type == "s_open":
            strike += 1
        elif child.type == "s_close":
            strike = max(0, strike - 1)

        elif child.type == "link_open":
            link_target = child.attrGet("href") or ""
        elif child.type == "link_close":
            link_target = None


def _render_table(
    document: DocumentType,
    tokens: list[Token],
    start_index: int,
    preset: StylePreset,
) -> int:
    """Render a Markdown table, preserving all inline formatting in every cell.

    Stores Token | None (inline tokens) rather than plain strings so that bold,
    italic, code, and links inside table cells are faithfully reproduced.
    """
    i = start_index + 1
    # header_inlines: inline token (or None) for each header column.
    header_inlines: list[Token | None] = []
    # row_inlines: list of rows; each row is a list of inline tokens (or None).
    row_inlines: list[list[Token | None]] = []
    current_row: list[Token | None] = []

    while i < len(tokens):
        token = tokens[i]
        if token.type == "table_close":
            break
        if token.type == "tr_open":
            current_row = []
        elif token.type == "tr_close":
            if header_inlines:
                row_inlines.append(current_row.copy())
            else:
                header_inlines = current_row.copy()
        elif token.type in {"th_open", "td_open"}:
            next_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            if next_tok and next_tok.type == "inline":
                current_row.append(next_tok)
            else:
                current_row.append(None)
        i += 1

    if not header_inlines:
        return i + 1

    col_count = max(len(header_inlines), max((len(r) for r in row_inlines), default=0))
    table = document.add_table(
        rows=max(1, len(row_inlines) + 1),
        cols=col_count,
        style="Table Grid",
    )
    table.autofit = True

    # Header row.
    for c, inline_token in enumerate(header_inlines):
        cell = table.cell(0, c)
        p = cell.paragraphs[0]
        if inline_token:
            _append_inline(p, inline_token, preset)
        _format_header_cell(cell, preset)

    # Body rows.
    for r, row in enumerate(row_inlines, start=1):
        for c, inline_token in enumerate(row):
            cell = table.cell(r, c)
            p = cell.paragraphs[0]
            if inline_token:
                _append_inline(p, inline_token, preset)
            _format_body_cell(cell)

    return i + 1


def _extract_report_header(markdown: str) -> ReportHeader:
    lines = markdown.splitlines()
    title = ""
    title_index: int | None = None
    for index, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            title_index = index
            break

    if title_index is None:
        return ReportHeader(title="", metadata=(), body_markdown=markdown)

    metadata: list[tuple[str, str]] = []
    end_index: int | None = None
    metadata_re = re.compile(r"^\*\*(.+?):\*\*\s*(.+?)\s*$")
    for index in range(title_index + 1, min(len(lines), title_index + 12)):
        stripped = lines[index].strip()
        if stripped == "---":
            end_index = index + 1
            break
        if not stripped:
            continue
        match = metadata_re.match(stripped)
        if not match:
            break
        metadata.append((match.group(1), match.group(2).strip()))

    if not metadata or end_index is None:
        return ReportHeader(title=title, metadata=(), body_markdown=markdown)

    body = "\n".join(lines[end_index:]).lstrip()
    return ReportHeader(title=title, metadata=tuple(metadata), body_markdown=body)


def _add_cover_page(document: DocumentType, header: ReportHeader, preset: StylePreset) -> None:
    # ── Top accent rule ───────────────────────────────────────────────────────
    accent_bar = document.add_paragraph()
    accent_bar.paragraph_format.space_before = Pt(0)
    accent_bar.paragraph_format.space_after = Pt(0)
    _set_paragraph_bottom_border(accent_bar, color=rgb_hex(preset.accent_color), size="24")

    # ── Title ─────────────────────────────────────────────────────────────────
    title = document.add_paragraph(style="Title")
    title.add_run(header.title)
    title.paragraph_format.space_before = Pt(60)
    title.paragraph_format.space_after = Pt(6)
    _set_paragraph_bottom_border(title, color=rgb_hex(preset.accent_color), size="12")

    # ── Metadata block ────────────────────────────────────────────────────────
    for label, value in header.metadata:
        para = document.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        label_run = para.add_run(f"{label}:")
        label_run.bold = True
        label_run.font.name = preset.body_font
        label_run.font.size = Pt(preset.body_size)
        # Tab-align the value column for a clean two-column appearance.
        para.add_run("\t")
        value_run = para.add_run(value)
        value_run.font.name = preset.body_font
        value_run.font.size = Pt(preset.body_size)

    document.add_page_break()


def _apply_report_header_footer(
    document: DocumentType,
    title: str,
    preset: StylePreset,
) -> None:
    for section in document.sections:
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = title
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.paragraph_format.space_after = Pt(0)
        for run in header_para.runs:
            run.font.name = preset.body_font
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(120, 120, 120)

        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.paragraph_format.space_before = Pt(0)
        run = footer_para.add_run("Page ")
        run.font.name = preset.body_font
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)
        _append_page_number(footer_para)


def _append_page_number(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instruction, end):
        run = paragraph.add_run()
        run._r.append(element)


def _append_preserved_text(run, text: str) -> None:
    lines = text.splitlines() or [""]
    for index, line in enumerate(lines):
        if index:
            run.add_break()
        run.add_text(line)


def _format_header_cell(cell, preset: StylePreset) -> None:
    _shade_cell(cell, preset.table_header_fill)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.bold = True


def _format_body_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)


def _shade_paragraph(paragraph, fill: str) -> None:
    """Apply a solid background fill to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _shade_cell(cell, fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _apply_run_shading(run, fill: str) -> None:
    """Apply a character-level background highlight to an inline run."""
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)


def _set_paragraph_bottom_border(paragraph, *, color: str, size: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _set_paragraph_left_border(paragraph, *, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "8")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)


# rgb_hex is provided by styles.py and re-exported from there.
